"""
Enhanced AI-Powered Resume Matcher v2.0
Integrates Google Cloud Vision API + Gemini AI + BM25 + FAISS
"""

import streamlit as st
import google.generativeai as genai
from google.cloud import vision
import io
import json
import tempfile
import os
import re
import traceback
import logging
from datetime import datetime

# NLP and ML libraries
import nltk
import spacy
from rank_bm25 import BM25Okapi
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

# Document processing
import PyPDF2
import docx

# Data processing and visualization
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    st.error("⚠️ Please install spaCy model: python -m spacy download en_core_web_sm")
    nlp = None

# ============================================================================
# CONFIGURATION
# ============================================================================

# Load environment variables (optional)
from dotenv import load_dotenv
load_dotenv()

# Gemini API Configuration
# For production, use environment variable: os.getenv("GEMINI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyD06chb5o4PMqghGspRqZVPsGBzEZ0S7vI")
genai.configure(api_key=GEMINI_API_KEY)

# Scoring weights for ensemble model
WEIGHTS = {
    "bm25": float(os.getenv("WEIGHT_BM25", 0.3)),      # Lexical matching
    "faiss": float(os.getenv("WEIGHT_FAISS", 0.3)),     # Semantic similarity
    "gemini": float(os.getenv("WEIGHT_GEMINI", 0.4))    # AI contextual analysis
}

# ============================================================================
# ENHANCED RESUME PROCESSOR CLASS
# ============================================================================

class EnhancedResumeProcessor:
    """Handles text extraction using Google Cloud Vision and traditional methods"""
    
    def __init__(self):
        try:
            self.vision_client = vision.ImageAnnotatorClient()
            self.vision_available = True
        except Exception as e:
            logger.warning(f"Vision API not available: {str(e)}")
            self.vision_available = False
        
        try:
            # Try gemini-pro first (most widely available)
            self.gemini_model = genai.GenerativeModel('gemini-pro')
            self.gemini_available = True
        except Exception as e:
            logger.warning(f"Gemini API not available: {str(e)}")
            self.gemini_available = False
    
    def extract_text(self, file):
        """Extract text from uploaded file using appropriate method"""
        file_type = file.type
        
        # For images, use Vision API if available
        if file_type in ["image/jpeg", "image/png", "image/jpg"]:
            if self.vision_available:
                return self._extract_image_vision(file)
            else:
                st.warning(f"⚠️ Vision API not configured. Cannot process image: {file.name}")
                return ""
        
        # For PDFs, try Vision API first for better OCR, then fallback
        elif file_type == "application/pdf":
            text = self._extract_pdf_traditional(file)
            return text
        
        # For DOCX files
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(file)
        
        else:
            st.warning(f"⚠️ Unsupported file type: {file_type}")
            return ""
    
    def _extract_image_vision(self, image_file):
        """Extract text from image using Google Cloud Vision API"""
        try:
            content = image_file.read()
            image = vision.Image(content=content)
            response = self.vision_client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(f"Vision API Error: {response.error.message}")
            
            if response.text_annotations:
                return response.text_annotations[0].description
            return ""
        except Exception as e:
            st.error(f"Error extracting text from image {image_file.name}: {str(e)}")
            return ""
    
    def _extract_pdf_traditional(self, pdf_file):
        """Extract text from PDF using PyPDF2"""
        try:
            text = ""
            reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Error reading page {page_num + 1}: {str(e)}")
                    continue
            
            return text.strip()
        except Exception as e:
            st.error(f"Error reading PDF {pdf_file.name}: {str(e)}")
            return ""
    
    def _extract_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            return "\n".join(paragraphs)
        except Exception as e:
            st.error(f"Error reading DOCX {docx_file.name}: {str(e)}")
            return ""
    
    def gemini_analyze_resume(self, job_description, resume_text, resume_name):
        """Use Gemini AI to analyze resume-job fit"""
        if not self.gemini_available:
            return None
        
        prompt = f"""You are an expert HR recruiter analyzing resume-job fit. Compare this job description with the candidate's resume:

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Provide detailed analysis in JSON format:
{{
    "match_percentage": <number 0-100>,
    "matching_skills": [<list of matching skills>],
    "missing_requirements": [<list of missing critical requirements>],
    "experience_alignment": "<brief assessment of experience match>",
    "strengths": [<list of candidate strengths>],
    "concerns": [<list of potential concerns>],
    "recommendation": "<hire/interview/reject with brief reasoning>"
}}

Be precise and professional. Return ONLY valid JSON, no additional text."""
        
        try:
            response = self.gemini_model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Clean up response to extract JSON
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            analysis = json.loads(response_text)
            return analysis
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error for {resume_name}: {str(e)}")
            logger.error(f"Response text: {response_text}")
            return None
        except Exception as e:
            logger.error(f"Gemini API error for {resume_name}: {str(e)}")
            return None

# ============================================================================
# TEXT PROCESSING FUNCTIONS
# ============================================================================

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove special characters but keep spaces
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    
    # Remove extra whitespace
    text = ' '.join(text.split())
    
    return text

def extract_skills(text):
    """Extract skills from resume text"""
    if not text or not nlp:
        return []
    
    # Common technical skills keywords
    skill_patterns = [
        r'\b(python|java|javascript|c\+\+|ruby|php|swift|kotlin|go|rust)\b',
        r'\b(react|angular|vue|node\.?js|django|flask|spring|express)\b',
        r'\b(sql|mysql|postgresql|mongodb|redis|elasticsearch)\b',
        r'\b(aws|azure|gcp|docker|kubernetes|jenkins|git)\b',
        r'\b(machine learning|deep learning|nlp|computer vision|data science)\b',
        r'\b(agile|scrum|devops|ci/cd|microservices|rest api)\b',
    ]
    
    skills = set()
    text_lower = text.lower()
    
    for pattern in skill_patterns:
        matches = re.findall(pattern, text_lower)
        skills.update(matches)
    
    # Use spaCy for additional entity extraction
    doc = nlp(text[:10000])  # Limit text length for performance
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PRODUCT", "LANGUAGE"]:
            skills.add(ent.text.lower())
    
    return list(skills)[:20]  # Limit to top 20 skills

def extract_experience_years(text):
    """Extract years of experience from resume"""
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s+experience',
    ]
    
    years = []
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        years.extend([int(y) for y in matches])
    
    return max(years) if years else 0

# ============================================================================
# SCORING FUNCTIONS
# ============================================================================

def build_bm25_index(resume_texts):
    """Build BM25 index from resume texts"""
    try:
        tokenized_corpus = []
        for text in resume_texts:
            tokens = nltk.word_tokenize(clean_text(text))
            tokenized_corpus.append(tokens)
        
        bm25 = BM25Okapi(tokenized_corpus)
        return bm25
    except Exception as e:
        st.error(f"Error building BM25 index: {str(e)}")
        return None

def build_faiss_index(resume_texts, resume_names):
    """Build FAISS vector index from resume texts"""
    try:
        model_name = "sentence-transformers/all-MiniLM-L6-v2"
        embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

        documents = []
        for text, name in zip(resume_texts, resume_names):
            doc = Document(
                page_content=text,
                metadata={"name": name}
            )
            documents.append(doc)

        vector_store = FAISS.from_documents(documents, embeddings)
        return vector_store
    except Exception as e:
        st.error(f"Error building FAISS index: {str(e)}")
        return None

def calculate_bm25_scores(bm25_index, job_description):
    """Calculate BM25 scores for all resumes"""
    jd_tokens = nltk.word_tokenize(clean_text(job_description))
    scores = bm25_index.get_scores(jd_tokens)

    # Normalize to 0-100 scale
    if len(scores) == 0:
        return [50.0]

    min_score, max_score = min(scores), max(scores)
    score_range = max(max_score - min_score, 0.001)

    normalized = [10 + 90 * ((score - min_score) / score_range) for score in scores]
    return normalized

def calculate_faiss_scores(vector_store, job_description, k):
    """Calculate FAISS semantic similarity scores"""
    try:
        results = vector_store.similarity_search_with_score(job_description, k=k)

        scores = []
        for doc, distance in results:
            # Convert distance to similarity percentage
            similarity = 100 * (1 - min(distance, 1.0))
            similarity = max(10, similarity)  # Minimum 10%
            scores.append({
                "name": doc.metadata["name"],
                "score": similarity
            })

        return scores
    except Exception as e:
        st.error(f"Error calculating FAISS scores: {str(e)}")
        return []

def calculate_ensemble_score(bm25_score, faiss_score, gemini_analysis):
    """Calculate final ensemble score using three models"""
    if not gemini_analysis:
        # Fallback to hybrid scoring without Gemini
        return round(0.6 * bm25_score + 0.4 * faiss_score, 2)

    gemini_score = gemini_analysis.get("match_percentage", 50)

    # Weighted ensemble
    final_score = (
        WEIGHTS["bm25"] * bm25_score +
        WEIGHTS["faiss"] * faiss_score +
        WEIGHTS["gemini"] * gemini_score
    )

    return round(final_score, 2)

# ============================================================================
# VISUALIZATION FUNCTIONS
# ============================================================================

def create_visualization(results):
    """Create bar chart visualization of results"""
    if not results:
        return

    fig, ax = plt.subplots(figsize=(12, 6))

    names = [r["name"][:30] for r in results]  # Truncate long names
    scores = [r["final_score"] for r in results]

    # Color coding based on score
    colors = []
    for score in scores:
        if score >= 70:
            colors.append('#2ecc71')  # Green
        elif score >= 50:
            colors.append('#f39c12')  # Orange
        else:
            colors.append('#e74c3c')  # Red

    bars = ax.barh(names, scores, color=colors)

    # Add score labels
    for i, (bar, score) in enumerate(zip(bars, scores)):
        ax.text(score + 1, i, f'{score:.1f}%', va='center', fontweight='bold')

    ax.set_xlabel('Match Score (%)', fontsize=12, fontweight='bold')
    ax.set_ylabel('Resume', fontsize=12, fontweight='bold')
    ax.set_title('Resume Match Scores', fontsize=14, fontweight='bold')
    ax.set_xlim(0, 110)

    plt.tight_layout()
    st.pyplot(fig)

def display_results(results):
    """Display detailed results for each resume"""
    st.subheader("📋 Detailed Results")

    for idx, result in enumerate(results, 1):
        # Color coding
        score = result["final_score"]
        if score >= 70:
            color = "🟢"
            status = "Strong Match"
        elif score >= 50:
            color = "🟡"
            status = "Moderate Match"
        else:
            color = "🔴"
            status = "Weak Match"

        with st.expander(f"{color} #{idx} - {result['name']} - {score}% ({status})"):
            col1, col2, col3 = st.columns(3)

            with col1:
                st.metric("Final Score", f"{result['final_score']}%")
                st.metric("BM25 Score", f"{result['bm25_score']:.1f}%")

            with col2:
                st.metric("FAISS Score", f"{result['faiss_score']:.1f}%")
                if result.get('gemini_score'):
                    st.metric("Gemini Score", f"{result['gemini_score']}%")

            with col3:
                st.metric("Skills Found", len(result.get('skills', [])))
                st.metric("Experience", f"{result.get('experience_years', 0)} years")

            # Gemini AI Analysis
            if result.get('gemini_analysis'):
                st.markdown("### 🤖 AI Analysis")
                analysis = result['gemini_analysis']

                if analysis.get('matching_skills'):
                    st.markdown("**✅ Matching Skills:**")
                    st.write(", ".join(analysis['matching_skills'][:10]))

                if analysis.get('missing_requirements'):
                    st.markdown("**❌ Missing Requirements:**")
                    st.write(", ".join(analysis['missing_requirements'][:5]))

                if analysis.get('strengths'):
                    st.markdown("**💪 Strengths:**")
                    for strength in analysis['strengths'][:3]:
                        st.write(f"- {strength}")

                if analysis.get('concerns'):
                    st.markdown("**⚠️ Concerns:**")
                    for concern in analysis['concerns'][:3]:
                        st.write(f"- {concern}")

                if analysis.get('recommendation'):
                    st.markdown(f"**📝 Recommendation:** {analysis['recommendation']}")

            # Skills
            if result.get('skills'):
                st.markdown("**🔧 Extracted Skills:**")
                st.write(", ".join(result['skills'][:15]))

def create_excel_download(results):
    """Create Excel file for download"""
    try:
        # Prepare data for Excel
        excel_data = []
        for idx, result in enumerate(results, 1):
            row = {
                "Rank": idx,
                "Resume Name": result["name"],
                "Final Score (%)": result["final_score"],
                "BM25 Score (%)": result["bm25_score"],
                "FAISS Score (%)": result["faiss_score"],
                "Gemini Score (%)": result.get("gemini_score", "N/A"),
                "Skills": ", ".join(result.get("skills", [])),
                "Experience (Years)": result.get("experience_years", 0),
            }

            # Add Gemini analysis if available
            if result.get('gemini_analysis'):
                analysis = result['gemini_analysis']
                row["AI Recommendation"] = analysis.get("recommendation", "")
                row["Matching Skills"] = ", ".join(analysis.get("matching_skills", [])[:10])
                row["Missing Requirements"] = ", ".join(analysis.get("missing_requirements", [])[:5])

            excel_data.append(row)

        df = pd.DataFrame(excel_data)

        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Resume Analysis')

        excel_bytes = output.getvalue()

        st.download_button(
            label="📥 Download Results (Excel)",
            data=excel_bytes,
            file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as e:
        st.error(f"Error creating Excel file: {str(e)}")

# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    st.set_page_config(
        page_title="Enhanced Resume Matcher v2.0",
        page_icon="🤖",
        layout="wide"
    )

    # Header
    st.title("🤖 AI-Powered Resume Matcher v2.0")
    st.markdown("### Enhanced with Google Cloud Vision + Gemini AI + BM25 + FAISS")

    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        st.markdown("**Scoring Weights:**")
        st.write(f"- BM25 (Lexical): {WEIGHTS['bm25']*100}%")
        st.write(f"- FAISS (Semantic): {WEIGHTS['faiss']*100}%")
        st.write(f"- Gemini AI: {WEIGHTS['gemini']*100}%")

        st.markdown("---")
        st.markdown("**Features:**")
        st.write("✅ Multi-format support (PDF, DOCX, Images)")
        st.write("✅ Google Cloud Vision OCR")
        st.write("✅ Gemini AI contextual analysis")
        st.write("✅ Hybrid BM25 + FAISS scoring")
        st.write("✅ Visual analytics")
        st.write("✅ Excel export")

        st.markdown("---")
        st.info("💡 Upload job description and resumes to start matching")

    # Initialize processor
    processor = EnhancedResumeProcessor()

    # Main content area
    col1, col2 = st.columns([2, 1])

    with col1:
        job_description = st.text_area(
            "📝 Job Description",
            height=250,
            placeholder="Paste the complete job description here...",
            help="Enter detailed job requirements, skills, and qualifications"
        )

    with col2:
        uploaded_files = st.file_uploader(
            "📎 Upload Resumes",
            type=["pdf", "docx", "jpg", "jpeg", "png"],
            accept_multiple_files=True,
            help="Supports PDF, DOCX, and image formats (JPG, PNG)"
        )

        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} file(s) uploaded")

    # Analysis button
    if st.button("🚀 Start Enhanced Analysis", type="primary", use_container_width=True):
        if not job_description:
            st.error("❌ Please enter a job description")
            return

        if not uploaded_files:
            st.error("❌ Please upload at least one resume")
            return

        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()

        # Step 1: Extract text from all resumes
        status_text.text("📄 Extracting text from resumes...")
        resume_data = []

        for idx, file in enumerate(uploaded_files):
            progress = (idx + 1) / len(uploaded_files) * 0.3
            progress_bar.progress(progress)

            text = processor.extract_text(file)

            if text:
                skills = extract_skills(text)
                experience_years = extract_experience_years(text)

                resume_data.append({
                    "name": file.name,
                    "text": text,
                    "skills": skills,
                    "experience_years": experience_years
                })

        if not resume_data:
            st.error("❌ No valid resumes processed. Please check your files.")
            return

        st.success(f"✅ Successfully processed {len(resume_data)} resumes")

        # Step 2: Build BM25 index
        status_text.text("🔍 Building BM25 index...")
        progress_bar.progress(0.4)

        resume_texts = [r["text"] for r in resume_data]
        resume_names = [r["name"] for r in resume_data]

        bm25_index = build_bm25_index(resume_texts)
        if not bm25_index:
            st.error("❌ Failed to build BM25 index")
            return

        # Step 3: Build FAISS index
        status_text.text("🧠 Building FAISS semantic index...")
        progress_bar.progress(0.5)

        vector_store = build_faiss_index(resume_texts, resume_names)
        if not vector_store:
            st.error("❌ Failed to build FAISS index")
            return

        # Step 4: Calculate scores
        status_text.text("📊 Calculating similarity scores...")
        progress_bar.progress(0.6)

        # BM25 scores
        bm25_scores = calculate_bm25_scores(bm25_index, job_description)

        # FAISS scores
        faiss_results = calculate_faiss_scores(vector_store, job_description, len(resume_data))
        faiss_score_map = {r["name"]: r["score"] for r in faiss_results}

        # Step 5: Gemini AI analysis
        status_text.text("🤖 Running Gemini AI analysis...")
        results = []

        for idx, resume in enumerate(resume_data):
            progress = 0.6 + (idx + 1) / len(resume_data) * 0.3
            progress_bar.progress(progress)

            # Get scores
            bm25_score = bm25_scores[idx]
            faiss_score = faiss_score_map.get(resume["name"], 50.0)

            # Gemini analysis
            gemini_analysis = processor.gemini_analyze_resume(
                job_description,
                resume["text"],
                resume["name"]
            )

            # Calculate ensemble score
            final_score = calculate_ensemble_score(bm25_score, faiss_score, gemini_analysis)

            results.append({
                "name": resume["name"],
                "final_score": final_score,
                "bm25_score": bm25_score,
                "faiss_score": faiss_score,
                "gemini_score": gemini_analysis.get("match_percentage") if gemini_analysis else None,
                "gemini_analysis": gemini_analysis,
                "skills": resume["skills"],
                "experience_years": resume["experience_years"]
            })

        # Sort by final score
        results.sort(key=lambda x: x["final_score"], reverse=True)

        progress_bar.progress(1.0)
        status_text.text("✅ Analysis complete!")

        # Display results
        st.markdown("---")
        display_results(results)

        # Visualization
        st.markdown("---")
        st.subheader("📊 Visual Analysis")
        create_visualization(results)

        # Excel download
        st.markdown("---")
        create_excel_download(results)

        # Summary statistics
        st.markdown("---")
        st.subheader("📈 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric("Total Resumes", len(results))

        with col2:
            avg_score = sum(r['final_score'] for r in results) / len(results)
            st.metric("Average Score", f"{avg_score:.1f}%")

        with col3:
            best_score = max(r['final_score'] for r in results)
            st.metric("Best Match", f"{best_score}%")

        with col4:
            total_skills = sum(len(r['skills']) for r in results)
            st.metric("Total Skills Found", total_skills)

if __name__ == "__main__":
    main()


