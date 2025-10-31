"""
Flask API Backend for Maharaj AI Resume Analyzer
Provides REST API endpoints for the frontend to interact with the resume analysis engine
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import tempfile
import traceback
from datetime import datetime

# Import the existing resume processing logic
import google.generativeai as genai
from google.cloud import vision
import nltk
import spacy
from rank_bm25 import BM25Okapi
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
import PyPDF2
import docx
import re
import json

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for frontend communication

# Configuration
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyD06chb5o4PMqghGspRqZVPsGBzEZ0S7vI")
genai.configure(api_key=GEMINI_API_KEY)

WEIGHTS = {
    "bm25": 0.3,
    "faiss": 0.3,
    "gemini": 0.4
}

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Warning: spaCy model not found. Run: python -m spacy download en_core_web_sm")
    nlp = None

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)


class EnhancedResumeProcessor:
    """Handles text extraction and AI analysis"""
    
    def __init__(self):
        try:
            self.vision_client = vision.ImageAnnotatorClient()
            self.vision_available = True
        except Exception as e:
            print(f"Vision API not available: {str(e)}")
            self.vision_available = False
        
        try:
            self.gemini_model = genai.GenerativeModel('gemini-pro')
            self.gemini_available = True
        except Exception as e:
            print(f"Gemini API not available: {str(e)}")
            self.gemini_available = False
    
    def extract_text(self, file_path, file_type):
        """Extract text from file"""
        if file_type in ["image/jpeg", "image/png", "image/jpg"]:
            if self.vision_available:
                return self._extract_image_vision(file_path)
        elif file_type == "application/pdf":
            return self._extract_pdf(file_path)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(file_path)
        return ""
    
    def _extract_image_vision(self, file_path):
        """Extract text from image using Vision API"""
        try:
            with open(file_path, 'rb') as image_file:
                content = image_file.read()
            image = vision.Image(content=content)
            response = self.vision_client.text_detection(image=image)
            
            if response.text_annotations:
                return response.text_annotations[0].description
            return ""
        except Exception as e:
            print(f"Error extracting text from image: {str(e)}")
            return ""
    
    def _extract_pdf(self, file_path):
        """Extract text from PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading PDF: {str(e)}")
            return ""
    
    def _extract_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"Error reading DOCX: {str(e)}")
            return ""
    
    def gemini_analyze_resume(self, job_description, resume_text, resume_name):
        """Use Gemini AI to analyze resume"""
        if not self.gemini_available:
            return None
        
        prompt = f"""You are an expert HR recruiter analyzing resume-job fit. Compare this job description with the candidate's resume:

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Provide detailed analysis in JSON format:
{{
    "match_percentage": <number 0-100>,
    "matching_skills": [<list of matching skills>],
    "missing_requirements": [<list of missing critical requirements>],
    "experience_alignment": "<brief assessment of experience match>",
    "strengths": [<list of candidate strengths>],
    "concerns": [<list of potential concerns>],
    "recommendation": "<hire/interview/reject with brief reasoning>"
}}

Be precise and professional. Return ONLY valid JSON, no additional text."""
        
        try:
            response = self.gemini_model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Clean up response
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            return json.loads(response_text)
        except Exception as e:
            print(f"Gemini API error: {str(e)}")
            return None


def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    text = ' '.join(text.split())
    return text


def extract_skills(text):
    """Extract skills from resume text"""
    if not text or not nlp:
        return []
    
    skill_patterns = [
        r'\b(python|java|javascript|c\+\+|ruby|php|swift|kotlin|go|rust)\b',
        r'\b(react|angular|vue|node\.?js|django|flask|spring|express)\b',
        r'\b(sql|mysql|postgresql|mongodb|redis|elasticsearch)\b',
        r'\b(aws|azure|gcp|docker|kubernetes|jenkins|git)\b',
        r'\b(machine learning|deep learning|nlp|computer vision|data science)\b',
        r'\b(agile|scrum|devops|ci/cd|microservices|rest api)\b',
    ]
    
    skills = set()
    text_lower = text.lower()
    
    for pattern in skill_patterns:
        matches = re.findall(pattern, text_lower)
        skills.update(matches)
    
    doc = nlp(text[:10000])
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PRODUCT", "LANGUAGE"]:
            skills.add(ent.text.lower())
    
    return list(skills)[:20]


def extract_experience_years(text):
    """Extract years of experience"""
    patterns = [
        r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yrs?\s+experience',
    ]
    
    years = []
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        years.extend([int(y) for y in matches])
    
    return max(years) if years else 0


def build_bm25_index(resume_texts):
    """Build BM25 index"""
    tokenized_corpus = []
    for text in resume_texts:
        tokens = nltk.word_tokenize(clean_text(text))
        tokenized_corpus.append(tokens)
    return BM25Okapi(tokenized_corpus)


def build_faiss_index(resume_texts, resume_names):
    """Build FAISS index"""
    model_name = "sentence-transformers/all-MiniLM-L6-v2"
    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    documents = [Document(page_content=text, metadata={"name": name}) 
                 for text, name in zip(resume_texts, resume_names)]
    
    return FAISS.from_documents(documents, embeddings)


def calculate_bm25_scores(bm25_index, job_description):
    """Calculate BM25 scores"""
    jd_tokens = nltk.word_tokenize(clean_text(job_description))
    scores = bm25_index.get_scores(jd_tokens)
    
    if len(scores) == 0:
        return [50.0]
    
    min_score, max_score = min(scores), max(scores)
    score_range = max(max_score - min_score, 0.001)
    
    return [10 + 90 * ((score - min_score) / score_range) for score in scores]


def calculate_faiss_scores(vector_store, job_description, k):
    """Calculate FAISS scores"""
    results = vector_store.similarity_search_with_score(job_description, k=k)
    scores = []
    for doc, distance in results:
        similarity = 100 * (1 - min(distance, 1.0))
        similarity = max(10, similarity)
        scores.append({"name": doc.metadata["name"], "score": similarity})
    return scores


def calculate_ensemble_score(bm25_score, faiss_score, gemini_analysis):
    """Calculate final ensemble score"""
    if not gemini_analysis:
        return round(0.6 * bm25_score + 0.4 * faiss_score, 2)

    gemini_score = gemini_analysis.get("match_percentage", 50)
    final_score = (
        WEIGHTS["bm25"] * bm25_score +
        WEIGHTS["faiss"] * faiss_score +
        WEIGHTS["gemini"] * gemini_score
    )
    return round(final_score, 2)


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'online',
        'timestamp': datetime.now().isoformat(),
        'version': '2.0'
    })


@app.route('/api/analyze', methods=['POST'])
def analyze_resumes():
    """Main endpoint for resume analysis"""
    try:
        # Get job description
        job_description = request.form.get('job_description')
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400

        # Get uploaded files
        files = request.files.getlist('resumes')
        if not files:
            return jsonify({'error': 'At least one resume is required'}), 400

        # Initialize processor
        processor = EnhancedResumeProcessor()

        # Process resumes
        resume_data = []
        temp_files = []

        for file in files:
            # Save file temporarily
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
            file.save(temp_file.name)
            temp_files.append(temp_file.name)

            # Extract text
            text = processor.extract_text(temp_file.name, file.content_type)

            if text:
                skills = extract_skills(text)
                experience_years = extract_experience_years(text)

                resume_data.append({
                    "name": file.filename,
                    "text": text,
                    "skills": skills,
                    "experience_years": experience_years
                })

        if not resume_data:
            return jsonify({'error': 'No valid resumes could be processed'}), 400

        # Build indices
        resume_texts = [r["text"] for r in resume_data]
        resume_names = [r["name"] for r in resume_data]

        bm25_index = build_bm25_index(resume_texts)
        vector_store = build_faiss_index(resume_texts, resume_names)

        # Calculate scores
        bm25_scores = calculate_bm25_scores(bm25_index, job_description)
        faiss_results = calculate_faiss_scores(vector_store, job_description, len(resume_data))
        faiss_score_map = {r["name"]: r["score"] for r in faiss_results}

        # Analyze with Gemini and create results
        results = []
        for idx, resume in enumerate(resume_data):
            bm25_score = bm25_scores[idx]
            faiss_score = faiss_score_map.get(resume["name"], 50.0)

            # Gemini analysis
            gemini_analysis = processor.gemini_analyze_resume(
                job_description,
                resume["text"],
                resume["name"]
            )

            # Calculate final score
            final_score = calculate_ensemble_score(bm25_score, faiss_score, gemini_analysis)

            results.append({
                "name": resume["name"],
                "final_score": final_score,
                "bm25_score": round(bm25_score, 1),
                "faiss_score": round(faiss_score, 1),
                "gemini_score": gemini_analysis.get("match_percentage") if gemini_analysis else None,
                "gemini_analysis": gemini_analysis,
                "skills": resume["skills"],
                "experience_years": resume["experience_years"]
            })

        # Sort by final score
        results.sort(key=lambda x: x["final_score"], reverse=True)

        # Update ranks
        for idx, result in enumerate(results):
            result["rank"] = idx + 1

        # Clean up temp files
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except:
                pass

        # Return results
        return jsonify({
            'success': True,
            'results': results,
            'summary': {
                'total_resumes': len(results),
                'average_score': round(sum(r['final_score'] for r in results) / len(results), 1),
                'best_match': results[0]['name'],
                'best_score': results[0]['final_score']
            }
        })

    except Exception as e:
        print(f"Error in analyze_resumes: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


@app.route('/api/config', methods=['GET'])
def get_config():
    """Get configuration settings"""
    return jsonify({
        'weights': WEIGHTS,
        'gemini_available': True,
        'vision_available': True
    })


if __name__ == '__main__':
    print("=" * 60)
    print("Maharaj AI Resume Analyzer - Backend API")
    print("=" * 60)
    print(f"Starting server on http://localhost:5000")
    print("API Endpoints:")
    print("  - GET  /api/health   - Health check")
    print("  - POST /api/analyze  - Analyze resumes")
    print("  - GET  /api/config   - Get configuration")
    print("=" * 60)

    app.run(debug=True, host='0.0.0.0', port=5000)

